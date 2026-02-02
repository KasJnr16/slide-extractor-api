"""
Advanced document generation services for creating Word, PDF, and text documents
with comprehensive styling, formatting, and packaging options.

Supports:
- Word documents (.docx) with full formatting
- PDF documents with text, tables, images
- Text documents (.txt, .md)
- ZIP packages with multiple documents
"""
import zipfile
from io import BytesIO
from typing import List, Dict, Any, Optional, Union
from datetime import datetime
from enum import Enum
import json
from .structured_logger import logger, OperationTimer
from .file_validator import validate_file


class DocumentFormat(Enum):
    """Supported document formats."""
    DOCX = "docx"
    PDF = "pdf"
    TXT = "txt"
    MARKDOWN = "md"


class DocumentStyle:
    """
    Document styling configuration.
    """
    def __init__(self,
                 font_family: str = "Arial",
                 font_size: int = 12,
                 heading1_size: int = 16,
                 heading2_size: int = 14,
                 heading3_size: int = 12,
                 line_spacing: float = 1.15,
                 margins: Optional[Dict[str, int]] = None,
                 page_size: str = "letter"):
        """
        Initialize document style.
        
        Args:
            font_family: Default font family
            font_size: Default font size in points
            heading1_size: H1 font size in points
            heading2_size: H2 font size in points
            heading3_size: H3 font size in points
            line_spacing: Line spacing multiplier
            margins: Dict with top, right, bottom, left in points
            page_size: "letter" or "a4"
        """
        self.font_family = font_family
        self.font_size = font_size
        self.heading1_size = heading1_size
        self.heading2_size = heading2_size
        self.heading3_size = heading3_size
        self.line_spacing = line_spacing
        self.margins = margins or {"top": 72, "right": 72, "bottom": 72, "left": 72}
        self.page_size = page_size
        
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "font_family": self.font_family,
            "font_size": self.font_size,
            "heading1_size": self.heading1_size,
            "heading2_size": self.heading2_size,
            "heading3_size": self.heading3_size,
            "line_spacing": self.line_spacing,
            "margins": self.margins,
            "page_size": self.page_size
        }


class ContentBlock:
    """Base class for document content blocks."""
    def __init__(self, block_type: str):
        self.block_type = block_type


class Paragraph(ContentBlock):
    """Text paragraph with formatting."""
    def __init__(self, text: str, 
                 bold: bool = False,
                 italic: bool = False,
                 underline: bool = False,
                 font_size: Optional[int] = None,
                 alignment: str = "left",
                 spacing_after: int = 6):
        """
        Create a paragraph.
        
        Args:
            text: Paragraph text
            bold: Bold text
            italic: Italic text
            underline: Underline text
            font_size: Font size (overrides default)
            alignment: "left", "center", "right", "justify"
            spacing_after: Space after paragraph in points
        """
        super().__init__("paragraph")
        self.text = text
        self.bold = bold
        self.italic = italic
        self.underline = underline
        self.font_size = font_size
        self.alignment = alignment
        self.spacing_after = spacing_after


class Heading(ContentBlock):
    """Heading with level."""
    def __init__(self, text: str, level: int = 1):
        """
        Create a heading.
        
        Args:
            text: Heading text
            level: Heading level (1-6)
        """
        super().__init__("heading")
        self.text = text
        self.level = max(1, min(6, level))  # Clamp between 1-6


class BulletList(ContentBlock):
    """Bulleted list."""
    def __init__(self, items: List[str], indent_level: int = 0):
        """
        Create a bullet list.
        
        Args:
            items: List of item texts
            indent_level: Indentation level (0-5)
        """
        super().__init__("bullet_list")
        self.items = items
        self.indent_level = indent_level


class NumberedList(ContentBlock):
    """Numbered list."""
    def __init__(self, items: List[str], indent_level: int = 0):
        """
        Create a numbered list.
        
        Args:
            items: List of item texts
            indent_level: Indentation level (0-5)
        """
        super().__init__("numbered_list")
        self.items = items
        self.indent_level = indent_level


class Table(ContentBlock):
    """Table with headers and data."""
    def __init__(self, headers: List[str], rows: List[List[str]],
                 column_widths: Optional[List[float]] = None,
                 border: bool = True,
                 header_background: Optional[str] = None):
        """
        Create a table.
        
        Args:
            headers: Column headers
            rows: Table data rows
            column_widths: Column widths as percentages (must sum to 100)
            border: Show borders
            header_background: Header background color (hex like "D5E8F0")
        """
        super().__init__("table")
        self.headers = headers
        self.rows = rows
        self.column_widths = column_widths
        self.border = border
        self.header_background = header_background or "D5E8F0"


class PageBreak(ContentBlock):
    """Page break."""
    def __init__(self):
        super().__init__("page_break")


class HorizontalLine(ContentBlock):
    """Horizontal line separator."""
    def __init__(self):
        super().__init__("horizontal_line")


class Image(ContentBlock):
    """Image with sizing."""
    def __init__(self, image_path: str, width: int, height: int,
                 alt_text: str = ""):
        """
        Create an image block.
        
        Args:
            image_path: Path to image file
            width: Width in points
            height: Height in points
            alt_text: Alt text for accessibility
        """
        super().__init__("image")
        self.image_path = image_path
        self.width = width
        self.height = height
        self.alt_text = alt_text


class DocumentGenerator:
    """
    Main document generator supporting Word, PDF, and text formats.
    """
    
    def __init__(self, style: Optional[DocumentStyle] = None):
        """
        Initialize generator.
        
        Args:
            style: Document style configuration
        """
        self.style = style or DocumentStyle()
    
    def create_document(self, title: str, content_blocks: List[ContentBlock], 
                        format: DocumentFormat = DocumentFormat.DOCX, 
                        request_id: str = None) -> bytes:
        """
        Create a document from content blocks.
        
        Args:
            title: Document title
            content_blocks: List of content blocks
            format: Output format
            request_id: Optional request ID for tracing
            
        Returns:
            Document as bytes
        """
        with OperationTimer("document_generation", title=title, format=format.value, request_id=request_id):
            try:
                logger.info(
                    "Starting document generation",
                    title=title,
                    format=format.value,
                    content_blocks=len(content_blocks),
                    request_id=request_id
                )
                
                if format == DocumentFormat.DOCX:
                    result = self._create_docx(title, content_blocks, metadata=None, request_id=request_id)
                elif format == DocumentFormat.PDF:
                    result = self._create_pdf(title, content_blocks, metadata=None, request_id=request_id)
                elif format == DocumentFormat.TXT:
                    result = self._create_txt(title, content_blocks, request_id=request_id)
                elif format == DocumentFormat.MARKDOWN:
                    result = self._create_markdown(title, content_blocks, request_id=request_id)
                else:
                    raise ValueError(f"Unsupported format: {format}")

                # Ensure bytes output for all formats
                if isinstance(result, str):
                    result = result.encode("utf-8")
                
                logger.log_file_processing(
                    filename=f"{title}.{format.value}",
                    file_size=len(result),
                    file_type="document",
                    operation="generation",
                    success=True,
                    request_id=request_id
                )
                
                return result
                
            except Exception as e:
                logger.error(
                    "Document generation failed",
                    title=title,
                    format=format.value,
                    error=str(e),
                    request_id=request_id
                )
                raise
    
    def _create_docx(self, title: str, content: List[ContentBlock],
                     metadata: Optional[Dict[str, str]] = None, request_id: str = None) -> bytes:
        """Create Word document using docx library."""
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # Set document metadata
        if metadata:
            core_props = doc.core_properties
            core_props.author = metadata.get('author', '')
            core_props.title = title
            core_props.subject = metadata.get('subject', '')
            core_props.comments = metadata.get('comments', '')
        
        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process content blocks
        for block in content:
            if isinstance(block, Paragraph):
                self._add_docx_paragraph(doc, block)
            
            elif isinstance(block, Heading):
                self._add_docx_heading(doc, block)
            
            elif isinstance(block, BulletList):
                self._add_docx_bullet_list(doc, block)
            
            elif isinstance(block, NumberedList):
                self._add_docx_numbered_list(doc, block)
            
            elif isinstance(block, Table):
                self._add_docx_table(doc, block)
            
            elif isinstance(block, PageBreak):
                doc.add_page_break()
            
            elif isinstance(block, HorizontalLine):
                self._add_docx_horizontal_line(doc)
            
            elif isinstance(block, Image):
                self._add_docx_image(doc, block)
        
        # Save to bytes
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    def _add_docx_paragraph(self, doc, para: Paragraph):
        """Add formatted paragraph to docx."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        p = doc.add_paragraph()
        run = p.add_run(para.text)
        
        # Apply formatting
        if para.bold:
            run.bold = True
        if para.italic:
            run.italic = True
        if para.underline:
            run.underline = True
        
        # Font size
        font_size = para.font_size or self.style.font_size
        run.font.size = Pt(font_size)
        run.font.name = self.style.font_family
        
        # Alignment
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        p.alignment = alignment_map.get(para.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # Spacing
        p.paragraph_format.space_after = Pt(para.spacing_after)
    
    def _add_docx_heading(self, doc, heading: Heading):
        """Add heading to docx."""
        h = doc.add_heading(heading.text, level=heading.level)
    
    def _add_docx_bullet_list(self, doc, bullet_list: BulletList):
        """Add bullet list to docx."""
        from docx.shared import Pt
        
        for item in bullet_list.items:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Pt(18 * bullet_list.indent_level)
    
    def _add_docx_numbered_list(self, doc, numbered_list: NumberedList):
        """Add numbered list to docx."""
        from docx.shared import Pt
        
        for item in numbered_list.items:
            p = doc.add_paragraph(item, style='List Number')
            p.paragraph_format.left_indent = Pt(18 * numbered_list.indent_level)
    
    def _add_docx_table(self, doc, table: Table):
        """Add table to docx."""
        from docx.shared import Pt, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Create table
        t = doc.add_table(rows=1, cols=len(table.headers))
        t.style = 'Light Grid Accent 1' if table.border else 'Table Grid'
        
        # Add headers
        header_cells = t.rows[0].cells
        for i, header in enumerate(table.headers):
            cell = header_cells[i]
            cell.text = header
            
            # Bold header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(self.style.font_size)
            
            # Background color
            if table.header_background:
                self._set_cell_background(cell, table.header_background)
        
        # Add data rows
        for row_data in table.rows:
            row_cells = t.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(self.style.font_size)
        
        # Add spacing after table
        doc.add_paragraph()
    
    def _set_cell_background(self, cell, color_hex: str):
        """Set cell background color in docx table."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def _add_docx_horizontal_line(self, doc):
        """Add horizontal line to docx."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 
                                   'w:kinsoku', 'w:wordWrap', 'w:overflowPunct')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
    
    def _add_docx_image(self, doc, image: Image):
        """Add image to docx."""
        from docx.shared import Inches
        import os
        
        if os.path.exists(image.image_path):
            doc.add_picture(image.image_path, 
                          width=Inches(image.width / 72),
                          height=Inches(image.height / 72))
            doc.add_paragraph()  # Spacing after image
    
    def _create_pdf(self, title: str, content: List[ContentBlock],
                    metadata: Optional[Dict[str, str]] = None, request_id: str = None) -> bytes:
        """Create PDF document using reportlab."""
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph as RLParagraph,
                                       Spacer, PageBreak as RLPageBreak, Table as RLTable,
                                       TableStyle, Image as RLImage)
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        
        # Setup
        output = BytesIO()
        page_size = letter if self.style.page_size == "letter" else A4
        
        doc = SimpleDocTemplate(
            output,
            pagesize=page_size,
            topMargin=self.style.margins['top'],
            rightMargin=self.style.margins['right'],
            bottomMargin=self.style.margins['bottom'],
            leftMargin=self.style.margins['left'],
            title=title,
            author=metadata.get('author', '') if metadata else ''
        )
        
        # Styles
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=self.style.heading1_size + 4,
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        # Add title
        story.append(RLParagraph(title, title_style))
        story.append(Spacer(1, 0.3 * inch))
        
        # Process content
        for block in content:
            if isinstance(block, Paragraph):
                story.extend(self._add_pdf_paragraph(block, styles))
            
            elif isinstance(block, Heading):
                story.extend(self._add_pdf_heading(block, styles))
            
            elif isinstance(block, BulletList):
                story.extend(self._add_pdf_bullet_list(block, styles))
            
            elif isinstance(block, NumberedList):
                story.extend(self._add_pdf_numbered_list(block, styles))
            
            elif isinstance(block, Table):
                story.extend(self._add_pdf_table(block))
            
            elif isinstance(block, PageBreak):
                story.append(RLPageBreak())
            
            elif isinstance(block, HorizontalLine):
                story.extend(self._add_pdf_horizontal_line())
            
            elif isinstance(block, Image):
                story.extend(self._add_pdf_image(block))
        
        # Build PDF
        doc.build(story)
        output.seek(0)
        return output.getvalue()
    
    def _add_pdf_paragraph(self, para: Paragraph, styles):
        """Add paragraph to PDF."""
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import Paragraph as RLParagraph, Spacer
        from reportlab.lib.units import inch
        
        # Create custom style
        alignment_map = {
            "left": TA_LEFT,
            "center": TA_CENTER,
            "right": TA_RIGHT,
            "justify": TA_JUSTIFY
        }
        
        style = ParagraphStyle(
            'CustomPara',
            parent=styles['Normal'],
            fontSize=para.font_size or self.style.font_size,
            alignment=alignment_map.get(para.alignment, TA_LEFT)
        )
        
        # Format text
        text = para.text
        if para.bold:
            text = f"<b>{text}</b>"
        if para.italic:
            text = f"<i>{text}</i>"
        if para.underline:
            text = f"<u>{text}</u>"
        
        return [
            RLParagraph(text, style),
            Spacer(1, para.spacing_after)
        ]
    
    def _add_pdf_heading(self, heading: Heading, styles):
        """Add heading to PDF."""
        from reportlab.platypus import Paragraph as RLParagraph, Spacer
        from reportlab.lib.units import inch
        
        style_name = f'Heading{heading.level}'
        return [
            RLParagraph(heading.text, styles[style_name]),
            Spacer(1, 0.1 * inch)
        ]
    
    def _add_pdf_bullet_list(self, bullet_list: BulletList, styles):
        """Add bullet list to PDF."""
        from reportlab.platypus import Paragraph as RLParagraph, Spacer
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch
        
        items = []
        style = ParagraphStyle(
            'BulletItem',
            parent=styles['Normal'],
            fontSize=self.style.font_size,
            leftIndent=bullet_list.indent_level * 20,
            bulletIndent=bullet_list.indent_level * 20 - 10
        )
        
        for item in bullet_list.items:
            items.append(RLParagraph(f"• {item}", style))
            items.append(Spacer(1, 3))
        
        return items
    
    def _add_pdf_numbered_list(self, numbered_list: NumberedList, styles):
        """Add numbered list to PDF."""
        from reportlab.platypus import Paragraph as RLParagraph, Spacer
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch
        
        items = []
        style = ParagraphStyle(
            'NumberedItem',
            parent=styles['Normal'],
            fontSize=self.style.font_size,
            leftIndent=numbered_list.indent_level * 20,
            bulletIndent=numbered_list.indent_level * 20 - 10
        )
        
        for i, item in enumerate(numbered_list.items, 1):
            items.append(RLParagraph(f"{i}. {item}", style))
            items.append(Spacer(1, 3))
        
        return items
    
    def _add_pdf_table(self, table: Table):
        """Add table to PDF."""
        from reportlab.platypus import Table as RLTable, TableStyle, Spacer
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        
        # Prepare data
        data = [table.headers] + table.rows
        
        # Create table
        t = RLTable(data, colWidths=table.column_widths)
        
        # Style
        style_commands = [
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), self.style.font_size),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]
        
        # Header background
        if table.header_background:
            # Convert hex to RGB
            hex_color = table.header_background.lstrip('#')
            r, g, b = tuple(int(hex_color[i:i+2], 16) / 255.0 for i in (0, 2, 4))
            style_commands.append(
                ('BACKGROUND', (0, 0), (-1, 0), colors.Color(r, g, b))
            )
        
        # Borders
        if table.border:
            style_commands.extend([
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BOX', (0, 0), (-1, -1), 1, colors.black),
            ])
        
        t.setStyle(TableStyle(style_commands))
        
        return [t, Spacer(1, 0.2 * inch)]
    
    def _add_pdf_horizontal_line(self):
        """Add horizontal line to PDF."""
        from reportlab.platypus import Spacer
        from reportlab.lib.units import inch
        from reportlab.platypus.flowables import HRFlowable
        from reportlab.lib import colors
        
        return [
            Spacer(1, 0.1 * inch),
            HRFlowable(width="100%", thickness=1, color=colors.grey),
            Spacer(1, 0.1 * inch)
        ]
    
    def _add_pdf_image(self, image: Image):
        """Add image to PDF."""
        from reportlab.platypus import Image as RLImage, Spacer
        from reportlab.lib.units import inch
        import os
        
        if os.path.exists(image.image_path):
            return [
                RLImage(image.image_path, 
                       width=image.width,
                       height=image.height),
                Spacer(1, 0.1 * inch)
            ]
        return []
    
    def _create_txt(self, title: str, content: List[ContentBlock], request_id: str = None) -> str:
        """Create plain text document."""
        lines = []
        lines.append("=" * 80)
        lines.append(title.center(80))
        lines.append("=" * 80)
        lines.append("")
        
        for block in content:
            if isinstance(block, Paragraph):
                lines.append(block.text)
                lines.append("")
            
            elif isinstance(block, Heading):
                lines.append("")
                lines.append("#" * block.level + " " + block.text)
                lines.append("-" * len(block.text))
                lines.append("")
            
            elif isinstance(block, BulletList):
                for item in block.items:
                    indent = "  " * block.indent_level
                    lines.append(f"{indent}• {item}")
                lines.append("")
            
            elif isinstance(block, NumberedList):
                for i, item in enumerate(block.items, 1):
                    indent = "  " * block.indent_level
                    lines.append(f"{indent}{i}. {item}")
                lines.append("")
            
            elif isinstance(block, Table):
                # Simple table formatting
                col_widths = [max(len(str(row[i])) for row in [block.headers] + block.rows)
                            for i in range(len(block.headers))]
                
                # Header
                header_row = " | ".join(block.headers[i].ljust(col_widths[i]) 
                                       for i in range(len(block.headers)))
                lines.append(header_row)
                lines.append("-" * len(header_row))
                
                # Rows
                for row in block.rows:
                    row_text = " | ".join(str(row[i]).ljust(col_widths[i]) 
                                        for i in range(len(row)))
                    lines.append(row_text)
                lines.append("")
            
            elif isinstance(block, PageBreak):
                lines.append("\n" * 3)
                lines.append("=" * 80)
                lines.append("\n" * 3)
            
            elif isinstance(block, HorizontalLine):
                lines.append("-" * 80)
                lines.append("")
        
        return "\n".join(lines)
    
    def _create_markdown(self, title: str, content: List[ContentBlock], request_id: str = None) -> str:
        """Create Markdown document."""
        lines = []
        lines.append(f"# {title}")
        lines.append("")
        
        for block in content:
            if isinstance(block, Paragraph):
                text = block.text
                if block.bold:
                    text = f"**{text}**"
                if block.italic:
                    text = f"*{text}*"
                lines.append(text)
                lines.append("")
            
            elif isinstance(block, Heading):
                lines.append(f"{'#' * (block.level + 1)} {block.text}")
                lines.append("")
            
            elif isinstance(block, BulletList):
                for item in block.items:
                    indent = "  " * block.indent_level
                    lines.append(f"{indent}- {item}")
                lines.append("")
            
            elif isinstance(block, NumberedList):
                for i, item in enumerate(block.items, 1):
                    indent = "  " * block.indent_level
                    lines.append(f"{indent}{i}. {item}")
                lines.append("")
            
            elif isinstance(block, Table):
                # Markdown table
                lines.append("| " + " | ".join(block.headers) + " |")
                lines.append("|" + "|".join([" --- "] * len(block.headers)) + "|")
                for row in block.rows:
                    lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
                lines.append("")
            
            elif isinstance(block, PageBreak):
                lines.append("\n---\n")
            
            elif isinstance(block, HorizontalLine):
                lines.append("---")
                lines.append("")
        
        return "\n".join(lines)


class DocumentPackager:
    """
    Create packages of multiple documents in ZIP format.
    """
    
    @staticmethod
    def create_package(documents: Dict[str, bytes], 
                      package_name: str = "documents") -> bytes:
        """
        Create a ZIP package with multiple documents.
        
        Args:
            documents: Dict mapping filenames to document bytes
            package_name: Base name for the package
            
        Returns:
            bytes: ZIP file as bytes
        """
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, mode='w', compression=zipfile.ZIP_DEFLATED) as zipf:
            for filename, content in documents.items():
                zipf.writestr(filename, content)
        
        zip_buffer.seek(0)
        return zip_buffer.getvalue()
    
    @staticmethod
    def create_exam_package(document_name: str,
                          questions: List[str],
                          answers: List[str],
                          format: DocumentFormat = DocumentFormat.DOCX,
                          style: Optional[DocumentStyle] = None) -> bytes:
        """
        Create a package with separate questions and answers documents.
        
        Args:
            document_name: Name for the document package
            questions: List of question strings
            answers: List of answer strings
            format: Output format for documents
            style: Document styling
            
        Returns:
            bytes: ZIP file as bytes
        """
        generator = DocumentGenerator(style)
        
        # Create questions document
        questions_content = [
            Heading("Questions", level=1)
        ]
        for i, q in enumerate(questions, 1):
            if q.strip():
                questions_content.append(Paragraph(f"Q{i}. {q}", spacing_after=12))
        
        questions_bytes = generator.create_document(
            f"{document_name} - Questions",
            questions_content,
            format
        )
        
        # Create answers document
        answers_content = [
            Heading("Answers", level=1)
        ]
        for i, a in enumerate(answers, 1):
            if a.strip():
                answers_content.append(Paragraph(f"A{i}. {a}", spacing_after=12))
        
        answers_bytes = generator.create_document(
            f"{document_name} - Answers",
            answers_content,
            format
        )
        
        # Package documents
        ext = format.value
        documents = {
            f"{document_name}_questions.{ext}": questions_bytes,
            f"{document_name}_answers.{ext}": answers_bytes
        }
        
        return DocumentPackager.create_package(documents, document_name)


# Convenience functions for backward compatibility
def create_exam_package_in_memory(document_name, questions, answers):
    """
    Legacy function - creates Word doc package.
    For new code, use DocumentPackager.create_exam_package()
    """
    return DocumentPackager.create_exam_package(
        document_name,
        questions,
        answers,
        DocumentFormat.DOCX
    )