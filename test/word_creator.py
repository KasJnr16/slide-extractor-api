from docx import Document
from docx.shared import Pt
import zipfile
import os

def create_exam_package(document_name, questions, answers, output_dir=None):
    """
    Creates two Word documents (questions and answers) and zips them.

    Args:
        document_name (str): Base name for files and zip.
        questions (list of str): List of exam questions.
        answers (list of str): List of corresponding answers.
        output_dir (str, optional): Directory to save files. Defaults to current directory.

    Returns:
        str: Path to the created zip file.
    """
    if output_dir is None:
        output_dir = os.getcwd()

    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # File paths
    questions_file = os.path.join(output_dir, f"{document_name}_questions.docx")
    answers_file = os.path.join(output_dir, f"{document_name}_answers.docx")
    zip_file = os.path.join(output_dir, f"{document_name}.zip")

    # ------------------
    # Create questions doc
    # ------------------
    doc_q = Document()
    doc_q.add_heading(f"{document_name} - Questions", level=1)

    for i, q in enumerate(questions, 1):
        para = doc_q.add_paragraph(f"{i}. {q}")
        para.paragraph_format.space_after = Pt(6)

    doc_q.save(questions_file)

    # ------------------
    # Create answers doc
    # ------------------
    doc_a = Document()
    doc_a.add_heading(f"{document_name} - Answers", level=1)

    for i, a in enumerate(answers, 1):
        para = doc_a.add_paragraph(f"{i}. {a}")
        para.paragraph_format.space_after = Pt(6)

    doc_a.save(answers_file)

    # ------------------
    # Zip the files
    # ------------------
    with zipfile.ZipFile(zip_file, 'w') as zipf:
        zipf.write(questions_file, os.path.basename(questions_file))
        zipf.write(answers_file, os.path.basename(answers_file))

    return zip_file

# ------------------
# Example usage
# ------------------
if __name__ == "__main__":
    questions = [
        "What is the CIA triad in computer security?",
        "Explain the difference between symmetric and asymmetric encryption."
    ]
    answers = [
        "Confidentiality, Integrity, Availability",
        "Symmetric uses same key; asymmetric uses public/private key pair"
    ]

    zip_path = create_exam_package("SecurityExam", questions, answers)
    print(f"✅ Exam package created: {zip_path}")
